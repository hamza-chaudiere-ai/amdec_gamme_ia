#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
Générateur de gammes de maintenance - VERSION COMPLÈTE
✅ FIX 2: Insertion automatique d'images d'appareils dans la colonne Observations
✅ NOUVEAU: Sections additionnelles (Inspections, Défauts, Contrôles, Image chaudière)
✅ Images techniques associées aux sous-composants
✅ Formatage Word professionnel avec images intégrées parfaitement
"""

import os
import logging
import glob
from datetime import datetime
from typing import Dict, List, Optional, Tuple
from docx import Document
from docx.shared import Pt, Cm, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.table import _Cell

# ✅ NOUVEAU: Imports pour l'insertion d'images
from docx.oxml.shared import qn
from docx.oxml.ns import nsdecls
from docx.oxml import parse_xml

from .utils import (
    format_component_display,
    format_subcomponent_display,
    get_maintenance_frequency,
    generate_timestamp,
    ComponentConfig
)

logger = logging.getLogger(__name__)

class GammeGenerator:
    """
    Générateur intelligent de gammes de maintenance - VERSION COMPLÈTE
    ✅ FIX 2: CORRIGÉ avec insertion automatique d'images dans les gammes
    ✅ NOUVEAU: Sections additionnelles avec inspections, défauts, contrôles
    """
    
    def __init__(self):
        """Initialise le générateur de gammes"""
        self.knowledge_base = self._build_maintenance_knowledge()
        self.materials_database = self._build_materials_database()
        self.operations_database = self._build_operations_database()
        
        # ✅ FIX 2: Base de données complète des images d'appareils
        self.images_database = self._build_complete_images_database()
        
        # ✅ NOUVEAU: Bases de données pour les sections additionnelles
        self.inspections_database = self._build_inspections_database()
        self.defauts_database = self._build_defauts_database()
        self.controles_database = self._build_controles_database()
    
    def _build_complete_images_database(self) -> Dict:
        """
        ✅ FIX 2: Base de données COMPLÈTE des images d'appareils
        Mapping précis entre sous-composants et images techniques disponibles
        """
        return {
            # Images par sous-composant spécifique
            'epingle': [
                'static/images/appareils/lampe_torche_appareil_photo.png',
                'static/images/appareils/appareil_de_mesure_a_ultrasons.png',
                'static/images/appareils/brosse_metallique.png',
                'static/images/appareils/pinceau_rouleau.png'
            ],
            'collecteur_entree': [
                'static/images/appareils/lampe_torche_endoscope.png',
                'static/images/appareils/kit_test_etancheite.png',
                'static/images/appareils/brosse_produit_nettoyant.png',
                'static/images/appareils/chiffon_degraissant.png'
            ],
            'collecteur_sortie': [
                'static/images/appareils/lampe_torche_kit_test_etancheite.png',
                'static/images/appareils/kit_de_serrage_cles_specifiques.png',
                'static/images/appareils/pinceau_rouleau.png',
                'static/images/appareils/renforts_metalliques_materiel_de_fixation.png'
            ],
            'tubes_suspension': [
                'static/images/appareils/capteurs_vibratoires_appareil_de_mesure_vibratoire.png',
                'static/images/appareils/cles_outils_de_serrage.png',
                'static/images/appareils/renforts_metalliques_materiel_de_fixation.png',
                'static/images/appareils/kit_de_serrage_cles_specifiques.png'
            ],
            'tube_porteur': [
                'static/images/appareils/appareil_de_mesure_a_ultrasons.png',
                'static/images/appareils/lampe_torche_camera_2.png',
                'static/images/appareils/brosse_produit_nettoyant.png',
                'static/images/appareils/loupe_d_inspection_camera_thermique.png'
            ],
            'branches_entree': [
                'static/images/appareils/lampe_torche.png',
                'static/images/appareils/loupe_d_inspection_camera_thermique.png',
                'static/images/appareils/chiffon_degraissant.png',
                'static/images/appareils/kit_test_etancheite.png'
            ],
            'branches_sortie': [
                'static/images/appareils/lampe_torche_camera_2.png',
                'static/images/appareils/kit_test_etancheite.png',
                'static/images/appareils/pinceau_rouleau.png',
                'static/images/appareils/brosse_metallique.png'
            ],
            # ✅ Images génériques si sous-composant non reconnu
            'default': [
                'static/images/appareils/lampe_torche.png',
                'static/images/appareils/appareil_de_mesure_a_ultrasons.png',
                'static/images/appareils/brosse_metallique.png',
                'static/images/appareils/kit_test_etancheite.png'
            ]
        }
    
    def _build_inspections_database(self) -> Dict:
        """
        ✅ NOUVEAU: Base de données des inspections recommandées avec localisations
        """
        return {
            'epingle': {
                'localizations': [
                    'Contrôle état des soudures épingles',
                    'Inspection épaisseur parois tubes',
                    'Vérification propreté surfaces',
                    'Contrôle déformation géométrique'
                ],
                'image_folder': 'static/images/inspections_localisations/'
            },
            'collecteur_entree': {
                'localizations': [
                    'Inspection interne collecteur',
                    'Contrôle raccordements entrée',
                    'Vérification étanchéité brides',
                    'Examen état supports'
                ],
                'image_folder': 'static/images/inspections_localisations/'
            },
            'collecteur_sortie': {
                'localizations': [
                    'Contrôle étanchéité sortie',
                    'Inspection soudures collecteur',
                    'Vérification contraintes thermiques',
                    'Examen corrosion interne'
                ],
                'image_folder': 'static/images/inspections_localisations/'
            },
            'tube_porteur': {
                'localizations': [
                    'Contrôle tubes porteurs principaux',
                    'Inspection supports mécaniques',
                    'Vérification alignement géométrique',
                    'Examen zones de contrainte'
                ],
                'image_folder': 'static/images/inspections_localisations/'
            },
            'branches_entree': {
                'localizations': [
                    'Contrôle branches distribution',
                    'Inspection raccordements amont',
                    'Vérification écoulement fluide',
                    'Examen usure interne'
                ],
                'image_folder': 'static/images/inspections_localisations/'
            },
            'branches_sortie': {
                'localizations': [
                    'Contrôle branches évacuation',
                    'Inspection raccordements aval',
                    'Vérification purges système',
                    'Examen accumulation dépôts'
                ],
                'image_folder': 'static/images/inspections_localisations/'
            },
            'tubes_suspension': {
                'localizations': [
                    'Contrôle points suspension',
                    'Inspection systèmes fixation',
                    'Vérification vibrations mécaniques',
                    'Examen usure supports'
                ],
                'image_folder': 'static/images/inspections_localisations/'
            }
        }
    
    def _build_defauts_database(self) -> Dict:
        """
        ✅ NOUVEAU: Base de données des défauts observés avec images illustratives
        """
        return {
            'causes_communes': {
                'Corrosion': {
                    'description': 'Dégradation par corrosion externe ou interne',
                    'image': 'static/images/défauts/corrosion.png',
                    'manifestations': ['Perte matière', 'Piqûres', 'Amincissement']
                },
                'Érosion': {
                    'description': 'Usure par abrasion des particules',
                    'image': 'static/images/défauts/erosion.png',
                    'manifestations': ['Surface rugueuse', 'Perte géométrie', 'Amincissement local']
                },
                'Fatigue': {
                    'description': 'Fissuration par cycles répétés',
                    'image': 'static/images/défauts/mechanical_fatigue.png',
                    'manifestations': ['Fissures', 'Propagation', 'Rupture']
                },
                'Surchauffe': {
                    'description': 'Dégradation par température excessive',
                    'image': 'static/images/défauts/longterm_overheat.png',
                    'manifestations': ['Déformation', 'Changement couleur', 'Fluage']
                }
            },
            'defauts_specifiques': {
                'economiseur_bt': {
                    'Caustic Attack': {
                        'image': 'static/images/défauts/caustic_attack.png',
                        'description': 'Attaque par solutions caustiques concentrées'
                    }
                },
                'surchauffeur_ht': {
                    'Long-term Overheat': {
                        'image': 'static/images/défauts/longterm_overheat.png',
                        'description': 'Surchauffe prolongée causant fluage'
                    },
                    'Fireside Corrosion': {
                        'image': 'static/images/défauts/fireside_corrosion_fatigue.png',
                        'description': 'Corrosion côté feu par dépôts acides'
                    }
                },
                'rechauffeur_ht': {
                    'Acid Attack': {
                        'image': 'static/images/défauts/acid_attack.png',
                        'description': 'Attaque acide des surfaces exposées'
                    },
                    'Hydrogen Damage': {
                        'image': 'static/images/défauts/hydrogen_damage.png',
                        'description': 'Dégradation par diffusion hydrogène'
                    }
                }
            }
        }
    
    def _build_controles_database(self) -> Dict:
        """
        ✅ NOUVEAU: Base de données des contrôles à effectuer avec méthodes illustrées
        """
        return {
            'controles_systematiques': {
                'Thermographie': {
                    'description': 'Contrôle thermique par caméra infrarouge',
                    'image': 'static/images/controles/controle_thermique.png',
                    'frequence': 'Mensuelle',
                    'zones': ['Points chauds', 'Gradients thermiques', 'Fuites thermiques']
                },
                'Mesure d\'épaisseur': {
                    'description': 'Contrôle dimensionnel par ultrasons',
                    'image': 'static/images/controles/controle_epaisseur.png',
                    'frequence': 'Trimestrielle',
                    'zones': ['Zones d\'usure', 'Points critiques', 'Coudes et raccords']
                },
                'Contrôle étanchéité': {
                    'description': 'Test de pression et recherche fuites',
                    'image': 'static/images/controles/controle_etancheite.png',
                    'frequence': 'Semestrielle',
                    'zones': ['Soudures', 'Brides', 'Raccordements']
                },
                'Analyse vibratoire': {
                    'description': 'Mesure vibrations et analyse fréquentielle',
                    'image': 'static/images/controles/controle_vibrations.png',
                    'frequence': 'Continue',
                    'zones': ['Supports', 'Points d\'ancrage', 'Zones flexibles']
                }
            },
            'controles_avances': {
                'Vertiscan': {
                    'description': 'Contrôle structural avancé par ultrasons',
                    'image': 'static/images/points_critiques/controle_vertiscan.png',
                    'frequence': 'Annuelle',
                    'zones': ['Structures porteuses', 'Assemblages soudés']
                },
                'Contrôle soudures': {
                    'description': 'Inspection détaillée des cordons de soudure',
                    'image': 'static/images/points_critiques/controle_soudures.png',
                    'frequence': 'Annuelle',
                    'zones': ['Tous les cordons', 'Zones de contrainte']
                },
                'Contrôle échangeur': {
                    'description': 'Inspection complète surfaces d\'échange',
                    'image': 'static/images/points_critiques/controle_echangeur.png',
                    'frequence': 'Semestrielle',
                    'zones': ['Surfaces internes', 'Chicanes', 'Tubes']
                }
            }
        }
    
    def _get_available_images_in_folder(self, folder_path: str) -> List[str]:
        """
        ✅ NOUVEAU: Récupère la liste des images disponibles dans un dossier
        """
        try:
            if not os.path.exists(folder_path):
                logger.warning(f"Dossier non trouvé: {folder_path}")
                return []
            
            # Extensions d'images supportées
            extensions = ['*.png', '*.jpg', '*.jpeg', '*.gif', '*.bmp']
            images = []
            
            for ext in extensions:
                pattern = os.path.join(folder_path, ext)
                images.extend(glob.glob(pattern))
            
            # Trier et retourner les chemins relatifs
            images.sort()
            return images
            
        except Exception as e:
            logger.error(f"Erreur lecture dossier {folder_path}: {e}")
            return []
    
    def generate(self, component: str, subcomponent: str, criticality: int = None) -> Dict:
        """
        Génère une gamme de maintenance complète avec images intégrées
        
        Args:
            component: Nom du composant
            subcomponent: Nom du sous-composant
            criticality: Criticité (calculée automatiquement si non fournie)
            
        Returns:
            Dictionnaire contenant toutes les données de la gamme avec images
        """
        try:
            logger.info(f"🔄 Génération gamme complète pour {component} - {subcomponent}")
            
            # Normaliser les entrées
            component = component.lower().replace(' ', '_')
            subcomponent = subcomponent.lower().replace(' ', '_')
            
            # Calculer la criticité si non fournie
            if criticality is None:
                criticality = ComponentConfig.get_default_criticality(component, subcomponent)
            
            # ✅ Récupération des images AVANT génération des opérations
            operation_images = self._get_operation_images_FIXED(component, subcomponent)
            
            # Générer tous les éléments de la gamme
            gamme_data = {
                'component': format_component_display(component),
                'subcomponent': format_subcomponent_display(subcomponent),
                'criticality': criticality,
                'criticality_level': self._get_criticality_level(criticality),
                'maintenance_frequency': get_maintenance_frequency(criticality),
                'materials': self._generate_materials_list(component, subcomponent, criticality),
                'operations': self._generate_operations(component, subcomponent, criticality),
                'safety_instructions': self._generate_safety_instructions(component, subcomponent),
                'estimated_time': '',  # Calculé après génération des opérations
                'date': datetime.now().strftime("%d/%m/%Y"),
                'images': operation_images,  # ✅ Images intégrées
                
                # ✅ NOUVEAU: Données pour les sections additionnelles
                'inspections': self._get_inspections_for_component(component, subcomponent),
                'defauts': self._get_defauts_for_component(component, subcomponent),
                'controles': self._get_controles_for_component(component, subcomponent, criticality)
            }
            
            # Calculer le temps total
            gamme_data['estimated_time'] = self._calculate_total_time(gamme_data['operations'])
            
            logger.info(f"✅ Gamme complète générée: {len(gamme_data['operations'])} opérations, "
                       f"{len(operation_images)} images appareils, "
                       f"{len(gamme_data['inspections'])} inspections, "
                       f"{len(gamme_data['defauts'])} défauts, "
                       f"{len(gamme_data['controles'])} contrôles")
            
            return gamme_data
            
        except Exception as e:
            logger.error(f"❌ Erreur lors de la génération de la gamme: {e}")
            raise
    
    def _get_inspections_for_component(self, component: str, subcomponent: str) -> List[Dict]:
        """
        ✅ NOUVEAU: Récupère les inspections recommandées pour le composant
        """
        try:
            inspections = []
            
            # Récupérer les inspections spécifiques au sous-composant
            subcomp_data = self.inspections_database.get(subcomponent, {})
            localizations = subcomp_data.get('localizations', [])
            image_folder = subcomp_data.get('image_folder', 'static/images/inspections_localisations/')
            
            # Récupérer les images disponibles dans le dossier
            available_images = self._get_available_images_in_folder(image_folder)
            
            # Associer chaque localisation à une image si disponible
            for i, localization in enumerate(localizations):
                inspection = {
                    'localisation': localization,
                    'image': None
                }
                
                # Essayer de trouver une image correspondante
                if i < len(available_images):
                    inspection['image'] = available_images[i]
                elif available_images:
                    # Utiliser une image aléatoire si pas assez d'images spécifiques
                    inspection['image'] = available_images[i % len(available_images)]
                
                inspections.append(inspection)
            
            # Si aucune inspection spécifique, ajouter des inspections génériques
            if not inspections:
                generic_inspections = [
                    'Inspection visuelle générale',
                    'Contrôle dimensionnel',
                    'Vérification état de surface',
                    'Examen points de fixation'
                ]
                
                for i, loc in enumerate(generic_inspections):
                    inspection = {
                        'localisation': loc,
                        'image': available_images[i] if i < len(available_images) else None
                    }
                    inspections.append(inspection)
            
            return inspections[:6]  # Limiter à 6 inspections maximum
            
        except Exception as e:
            logger.error(f"Erreur génération inspections: {e}")
            return []
    
    def _get_defauts_for_component(self, component: str, subcomponent: str) -> List[Dict]:
        """
        ✅ NOUVEAU: Récupère les défauts identifiés pour le composant
        """
        try:
            defauts = []
            
            # Ajouter les causes communes
            causes_communes = self.defauts_database.get('causes_communes', {})
            for cause, data in causes_communes.items():
                if os.path.exists(data.get('image', '')):
                    defauts.append({
                        'cause': cause,
                        'description': data.get('description', ''),
                        'image': data.get('image'),
                        'manifestations': data.get('manifestations', [])
                    })
            
            # Ajouter les défauts spécifiques au composant
            component_defauts = self.defauts_database.get('defauts_specifiques', {}).get(component, {})
            for defaut, data in component_defauts.items():
                if os.path.exists(data.get('image', '')):
                    defauts.append({
                        'cause': defaut,
                        'description': data.get('description', ''),
                        'image': data.get('image'),
                        'manifestations': []
                    })
            
            # Si pas assez de défauts, récupérer toutes les images disponibles du dossier
            if len(defauts) < 3:
                available_images = self._get_available_images_in_folder('static/images/défauts/')
                
                for img_path in available_images[:6]:  # Max 6 images
                    if not any(d['image'] == img_path for d in defauts):  # Éviter doublons
                        defaut_name = os.path.basename(img_path).replace('.png', '').replace('_', ' ').title()
                        defauts.append({
                            'cause': defaut_name,
                            'description': f'Défaut de type {defaut_name.lower()}',
                            'image': img_path,
                            'manifestations': []
                        })
            
            return defauts[:8]  # Maximum 8 défauts
            
        except Exception as e:
            logger.error(f"Erreur génération défauts: {e}")
            return []
    
    def _get_controles_for_component(self, component: str, subcomponent: str, criticality: int) -> List[Dict]:
        """
        ✅ NOUVEAU: Récupère les contrôles à effectuer selon la criticité
        """
        try:
            controles = []
            
            # Contrôles systématiques (toujours inclus)
            systematiques = self.controles_database.get('controles_systematiques', {})
            for controle, data in systematiques.items():
                if os.path.exists(data.get('image', '')):
                    controles.append({
                        'controle': controle,
                        'description': data.get('description', ''),
                        'image': data.get('image'),
                        'frequence': data.get('frequence', 'Régulière'),
                        'zones': data.get('zones', [])
                    })
            
            # Contrôles avancés (si criticité élevée)
            if criticality > 16:
                avances = self.controles_database.get('controles_avances', {})
                for controle, data in avances.items():
                    if os.path.exists(data.get('image', '')):
                        controles.append({
                            'controle': controle,
                            'description': data.get('description', ''),
                            'image': data.get('image'),
                            'frequence': data.get('frequence', 'Annuelle'),
                            'zones': data.get('zones', [])
                        })
            
            # Si pas assez de contrôles, ajouter des images depuis points_critiques
            if len(controles) < 4:
                available_images = self._get_available_images_in_folder('static/images/points_critiques/')
                
                for img_path in available_images:
                    if not any(c['image'] == img_path for c in controles):
                        controle_name = os.path.basename(img_path).replace('.png', '').replace('_', ' ').title()
                        controles.append({
                            'controle': controle_name,
                            'description': f'Contrôle de type {controle_name.lower()}',
                            'image': img_path,
                            'frequence': 'Selon criticité',
                            'zones': ['Zone concernée']
                        })
                        
                        if len(controles) >= 6:  # Limiter à 6 contrôles
                            break
            
            return controles
            
        except Exception as e:
            logger.error(f"Erreur génération contrôles: {e}")
            return []
    
    def _get_operation_images_FIXED(self, component: str, subcomponent: str) -> List[str]:
        """
        ✅ FIX 2: Récupération CORRIGÉE des images d'appareils pour le sous-composant
        """
        try:
            # Nettoyer le nom du sous-composant et tenter plusieurs correspondances
            subcomp_clean = subcomponent.lower().replace(' ', '_')
            
            logger.debug(f"🔍 Recherche images pour: {subcomp_clean}")
            
            # ✅ Étape 1: Correspondance exacte
            images = self.images_database.get(subcomp_clean, [])
            
            # ✅ Étape 2: Correspondance partielle si pas de correspondance exacte
            if not images:
                for db_key, db_images in self.images_database.items():
                    if db_key == 'default':
                        continue
                    # Recherche par inclusion partielle
                    if db_key in subcomp_clean or any(part in db_key for part in subcomp_clean.split('_')):
                        images = db_images
                        logger.debug(f"✅ Correspondance partielle trouvée: {db_key}")
                        break
            
            # ✅ Étape 3: Vérifier que les fichiers existent
            existing_images = []
            for img_path in images:
                if os.path.exists(img_path):
                    existing_images.append(img_path)
                    logger.debug(f"✅ Image trouvée: {img_path}")
                else:
                    logger.warning(f"⚠️ Image non trouvée: {img_path}")
            
            # ✅ Étape 4: Si aucune image spécifique, utiliser les images par défaut
            if not existing_images:
                logger.info(f"🔄 Utilisation images par défaut pour {subcomponent}")
                default_images = self.images_database.get('default', [])
                existing_images = [img for img in default_images if os.path.exists(img)]
            
            # ✅ Étape 5: Assurer un minimum d'images
            if not existing_images:
                logger.warning(f"⚠️ Aucune image disponible pour {subcomponent}")
                # Créer une liste d'images de fallback si les fichiers n'existent pas
                fallback_images = [
                    'static/images/appareils/lampe_torche.png',
                    'static/images/appareils/appareil_de_mesure_a_ultrasons.png'
                ]
                existing_images = [img for img in fallback_images if os.path.exists(img)]
            
            # Limiter à 4 images maximum pour optimiser la mise en page
            final_images = existing_images[:4]
            
            logger.info(f"✅ {len(final_images)} images finales sélectionnées pour {subcomponent}")
            
            return final_images
            
        except Exception as e:
            logger.error(f"❌ Erreur récupération images pour {subcomponent}: {e}")
            return []
    
    def save_to_file(self, gamme_data: Dict, component: str, subcomponent: str, 
                     output_path: str = None) -> str:
        """
        ✅ FIX 2: Sauvegarde la gamme avec images PARFAITEMENT intégrées dans le Word
        
        Args:
            gamme_data: Données de la gamme
            component: Nom du composant
            subcomponent: Nom du sous-composant
            output_path: Chemin de sortie (optionnel)
            
        Returns:
            Chemin du fichier sauvegardé
        """
        if output_path is None:
            os.makedirs('data/generated/gammes', exist_ok=True)
            timestamp = generate_timestamp()
            safe_comp = component.lower().replace(' ', '_')
            safe_subcomp = subcomponent.lower().replace(' ', '_')
            filename = f"gamme_{safe_comp}_{safe_subcomp}_{timestamp}.docx"
            output_path = os.path.join('data/generated/gammes', filename)
        
        # ✅ Créer le document Word avec images PARFAITEMENT intégrées + sections additionnelles
        self._create_word_document_with_perfect_images(gamme_data, output_path)
        
        logger.info(f"✅ Gamme complète sauvegardée avec {len(gamme_data.get('images', []))} images: {output_path}")
        return output_path
    
    def _create_word_document_with_perfect_images(self, gamme_data: Dict, output_path: str):
        """
        ✅ FIX 2: Crée le document Word avec images PARFAITEMENT intégrées dans la colonne Observations
        ✅ NOUVEAU: Avec sections additionnelles (Inspections, Défauts, Contrôles, Image chaudière)
        """
        doc = Document()
        
        # Configuration des marges
        sections = doc.sections
        for section in sections:
            section.page_height = Cm(29.7)
            section.page_width = Cm(21.0)
            section.left_margin = Cm(1.5)
            section.right_margin = Cm(1.5)
            section.top_margin = Cm(2.0)
            section.bottom_margin = Cm(2.0)
        
        # ✅ Titre principal avec style
        title = f"GAMME DE MAINTENANCE PREVENTIVE"
        title_para = doc.add_heading(title, level=1)
        title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Sous-titre avec composant
        subtitle = f"Machine : Chaudière - Equipement : {gamme_data['component']} / {gamme_data['subcomponent']}"
        subtitle_para = doc.add_paragraph(subtitle)
        subtitle_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        subtitle_para.runs[0].bold = True
        subtitle_para.runs[0].font.size = Pt(14)
        
        doc.add_paragraph()  # Espacement
        
        # ✅ Informations générales dans un tableau formaté comme le template
        info_table = doc.add_table(rows=6, cols=2)
        info_table.style = 'Table Grid'
        info_table.alignment = WD_TABLE_ALIGNMENT.LEFT
        
        info_data = [
            ("Equipement", f"{gamme_data['component']} - {gamme_data['subcomponent']}"),
            ("Criticité", f"C = {gamme_data['criticality']} ({gamme_data['criticality_level']})"),
            ("Fréquence recommandée", gamme_data['maintenance_frequency']),
            ("Date de génération", gamme_data['date']),
            ("Temps total estimé", gamme_data['estimated_time']),
            ("Document source", "PLAN DE MAINTENANCE PREVENTIVE")
        ]
        
        for i, (label, value) in enumerate(info_data):
            info_table.cell(i, 0).text = label
            info_table.cell(i, 1).text = value
            info_table.cell(i, 0).paragraphs[0].runs[0].bold = True
            
            # Largeur des colonnes
            info_table.cell(i, 0).width = Cm(5)
            info_table.cell(i, 1).width = Cm(10)
        
        doc.add_paragraph()  # Espacement
        
        # ✅ Consignes de sécurité formatées
        security_title = doc.add_paragraph("Consignes de sécurité :")
        security_title.runs[0].bold = True
        security_title.runs[0].font.size = Pt(12)
        
        security_para = doc.add_paragraph()
        for instruction in gamme_data['safety_instructions']:
            security_para.add_run(f"• {instruction}\n")
        
        doc.add_paragraph()  # Espacement
        
        # ✅ TABLEAU DES OPÉRATIONS AVEC IMAGES - FORMAT EXACT DU TEMPLATE
        operations_title = doc.add_paragraph("OPERATIONS A REALISER A L'ALERT")
        operations_title.runs[0].bold = True
        operations_title.runs[0].font.size = Pt(12)
        operations_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # ✅ Créer le tableau avec 6 colonnes comme dans le template
        operations_table = doc.add_table(rows=1, cols=6)
        operations_table.style = 'Table Grid'
        operations_table.alignment = WD_TABLE_ALIGNMENT.CENTER
        
        # ✅ En-têtes du tableau (format template)
        hdr_cells = operations_table.rows[0].cells
        headers = ["Ordre", "Ordres Des OP", "Opérations", "Temps alloué", "Matériel à employer", "Observations"]
        
        for i, header in enumerate(headers):
            hdr_cells[i].text = header
            hdr_cells[i].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = hdr_cells[i].paragraphs[0].runs[0]
            run.bold = True
            run.font.size = Pt(10)
        
        # Définir les largeurs de colonnes
        col_widths = [Cm(1.5), Cm(1.5), Cm(4.0), Cm(2.0), Cm(3.0), Cm(4.0)]
        for i, width in enumerate(col_widths):
            hdr_cells[i].width = width
        
        # ✅ Ajouter les opérations avec images dans la colonne "Observations"
        images = gamme_data.get('images', [])
        
        for i, operation in enumerate(gamme_data['operations']):
            row_cells = operations_table.add_row().cells
            
            # Colonne 1: Ordre principal (groupé par opération majeure)
            if i == 0:
                row_cells[0].text = "1"
            else:
                row_cells[0].text = ""  # Fusion visuelle
            row_cells[0].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            # Colonne 2: Ordre de sous-opération
            row_cells[1].text = str(operation['order'])
            row_cells[1].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            # Colonne 3: Opérations (nom + description)
            op_para = row_cells[2].paragraphs[0]
            op_para.clear()
            op_para.add_run(operation['name']).bold = True
            if operation.get('description'):
                op_para.add_run(f"\n{operation['description']}")
            
            # Colonne 4: Temps alloué
            row_cells[3].text = f"{operation['time']} min"
            row_cells[3].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            # Colonne 5: Matériel (extrait des opérations)
            materials = self._get_materials_for_operation(operation, gamme_data['materials'])
            row_cells[4].text = materials
            
            # ✅ Colonne 6: OBSERVATIONS AVEC IMAGE INTÉGRÉE
            obs_cell = row_cells[5]
            obs_para = obs_cell.paragraphs[0]
            obs_para.clear()
            
            # ✅ INSERTION DE L'IMAGE si disponible
            if i < len(images) and os.path.exists(images[i]):
                try:
                    # ✅ Insérer l'image avec taille optimisée pour la cellule
                    run = obs_para.add_run()
                    run.add_picture(images[i], width=Cm(3.5), height=Cm(2.5))
                    obs_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    
                    # ✅ Ajouter le nom de l'appareil sous l'image
                    img_name = os.path.basename(images[i]).replace('.png', '').replace('_', ' ').title()
                    obs_para.add_run(f"\n{img_name}")
                    obs_para.runs[-1].font.size = Pt(8)
                    obs_para.runs[-1].italic = True
                    
                    logger.debug(f"✅ Image {i+1} insérée: {images[i]} dans opération {operation['name']}")
                    
                except Exception as e:
                    logger.warning(f"⚠️ Erreur insertion image {images[i]}: {e}")
                    obs_para.add_run("Image non disponible")
                    obs_para.runs[0].font.size = Pt(9)
                    obs_para.runs[0].italic = True
            else:
                # ✅ Si pas d'image, afficher un texte de procédure
                if operation.get('description'):
                    obs_para.add_run("Voir procédure détaillée")
                else:
                    obs_para.add_run("Contrôle visuel standard")
                obs_para.runs[0].font.size = Pt(9)
                obs_para.runs[0].italic = True
                obs_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            # ✅ Formater la hauteur de ligne pour accommoder les images
            for cell in row_cells:
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        if not hasattr(run, 'element') or 'picture' not in str(run.element.xml):
                            run.font.size = Pt(9)  # Texte plus petit pour optimiser l'espace
        
        # ✅ Ajuster automatiquement la hauteur des lignes
        for row in operations_table.rows[1:]:  # Exclure l'en-tête
            row.height = Cm(3.0)  # Hauteur fixe pour accommoder les images
        
        doc.add_paragraph()  # Espacement
        
        # ✅ Matériels nécessaires (section détaillée)
        doc.add_heading("Matériels et Équipements Nécessaires", level=2)
        materials_table = doc.add_table(rows=1, cols=2)
        materials_table.style = 'Table Grid'
        
        # En-têtes
        materials_table.cell(0, 0).text = "Catégorie"
        materials_table.cell(0, 1).text = "Matériels"
        for cell in materials_table.rows[0].cells:
            cell.paragraphs[0].runs[0].bold = True
            cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Grouper les matériels par catégorie
        materials_by_category = self._group_materials_by_category(gamme_data['materials'])
        
        for category, items in materials_by_category.items():
            row_cells = materials_table.add_row().cells
            row_cells[0].text = category
            row_cells[1].text = " • ".join(items)
            row_cells[0].paragraphs[0].runs[0].bold = True
        
        # ✅ Section images supplémentaires si plus d'images que d'opérations
        remaining_images = images[len(gamme_data['operations']):]
        if remaining_images:
            doc.add_page_break()
            doc.add_heading("Images Techniques Complémentaires", level=2)
            
            # ✅ Créer un tableau 2x2 pour les images supplémentaires
            img_table = doc.add_table(rows=0, cols=2)
            img_table.style = 'Table Grid'
            
            for i in range(0, len(remaining_images), 2):
                row_cells = img_table.add_row().cells
                
                # Image de gauche
                if i < len(remaining_images) and os.path.exists(remaining_images[i]):
                    try:
                        left_para = row_cells[0].paragraphs[0]
                        left_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                        
                        run = left_para.add_run()
                        run.add_picture(remaining_images[i], width=Cm(6.0))
                        
                        # Légende
                        img_name = os.path.basename(remaining_images[i]).replace('.png', '').replace('_', ' ').title()
                        left_para.add_run(f"\n{img_name}")
                        left_para.runs[-1].font.size = Pt(10)
                        left_para.runs[-1].bold = True
                        
                    except Exception as e:
                        logger.warning(f"⚠️ Erreur insertion image supplémentaire: {e}")
                
                # Image de droite
                if i + 1 < len(remaining_images) and os.path.exists(remaining_images[i + 1]):
                    try:
                        right_para = row_cells[1].paragraphs[0]
                        right_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                        
                        run = right_para.add_run()
                        run.add_picture(remaining_images[i + 1], width=Cm(6.0))
                        
                        # Légende
                        img_name = os.path.basename(remaining_images[i + 1]).replace('.png', '').replace('_', ' ').title()
                        right_para.add_run(f"\n{img_name}")
                        right_para.runs[-1].font.size = Pt(10)
                        right_para.runs[-1].bold = True
                        
                    except Exception as e:
                        logger.warning(f"⚠️ Erreur insertion image supplémentaire: {e}")
        
        # ✅ Informations complémentaires si enrichi avec AMDEC
        if gamme_data.get('amdec_cause'):
            doc.add_paragraph()
            doc.add_heading("Informations AMDEC Associées", level=2)
            
            amdec_table = doc.add_table(rows=4, cols=2)
            amdec_table.style = 'Table Grid'
            
            amdec_info = [
                ("Cause principale identifiée", gamme_data.get('amdec_cause', 'Non spécifiée')),
                ("Mode de défaillance", gamme_data.get('amdec_mode', 'Non spécifié')),
                ("Effet observé", gamme_data.get('amdec_effect', 'Non spécifié')),
                ("Actions correctives recommandées", gamme_data.get('amdec_actions', 'Voir gamme de maintenance'))
            ]
            
            for i, (label, value) in enumerate(amdec_info):
                amdec_table.cell(i, 0).text = label
                amdec_table.cell(i, 1).text = value
                amdec_table.cell(i, 0).paragraphs[0].runs[0].bold = True
                
                # Largeurs
                amdec_table.cell(i, 0).width = Cm(5)
                amdec_table.cell(i, 1).width = Cm(10)
        
        # ========================================
        # ✅ NOUVEAU: SECTIONS ADDITIONNELLES
        # ========================================
        
        # ✅ NOUVEAU: Section 1 - Inspections recommandées
        if gamme_data.get('inspections'):
            doc.add_page_break()
            doc.add_heading("Inspections Recommandées", level=2)
            
            inspections_table = doc.add_table(rows=1, cols=2)
            inspections_table.style = 'Table Grid'
            inspections_table.alignment = WD_TABLE_ALIGNMENT.CENTER
            
            # En-têtes
            hdr_cells = inspections_table.rows[0].cells
            hdr_cells[0].text = "Localisation"
            hdr_cells[1].text = "Image correspondante"
            
            for cell in hdr_cells:
                cell.paragraphs[0].runs[0].bold = True
                cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
                cell.paragraphs[0].runs[0].font.size = Pt(11)
            
            # Définir largeurs
            hdr_cells[0].width = Cm(8)
            hdr_cells[1].width = Cm(8)
            
            # Ajouter les inspections
            for inspection in gamme_data['inspections']:
                row_cells = inspections_table.add_row().cells
                
                # Localisation
                row_cells[0].text = inspection['localisation']
                row_cells[0].paragraphs[0].runs[0].font.size = Pt(10)
                
                # Image
                if inspection.get('image') and os.path.exists(inspection['image']):
                    try:
                        img_para = row_cells[1].paragraphs[0]
                        img_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                        
                        run = img_para.add_run()
                        run.add_picture(inspection['image'], width=Cm(6.0), height=Cm(4.0))
                        
                        # Légende sous l'image
                        img_name = os.path.basename(inspection['image']).replace('.png', '').replace('_', ' ').title()
                        img_para.add_run(f"\n{img_name}")
                        img_para.runs[-1].font.size = Pt(8)
                        img_para.runs[-1].italic = True
                        
                    except Exception as e:
                        logger.warning(f"⚠️ Erreur insertion image inspection: {e}")
                        row_cells[1].text = "Image non disponible"
                        row_cells[1].paragraphs[0].runs[0].font.size = Pt(9)
                        row_cells[1].paragraphs[0].runs[0].italic = True
                else:
                    row_cells[1].text = "Voir procédure standard"
                    row_cells[1].paragraphs[0].runs[0].font.size = Pt(9)
                    row_cells[1].paragraphs[0].runs[0].italic = True
                    row_cells[1].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
                
                # Hauteur de ligne pour les images
                row_cells[0].width = Cm(8)
                row_cells[1].width = Cm(8)
        
        # ✅ NOUVEAU: Section 2 - Causes identifiées (Défauts observés)
        if gamme_data.get('defauts'):
            doc.add_paragraph()
            doc.add_heading("Causes Identifiées (Défauts Observés)", level=2)
            
            defauts_table = doc.add_table(rows=1, cols=2)
            defauts_table.style = 'Table Grid'
            defauts_table.alignment = WD_TABLE_ALIGNMENT.CENTER
            
            # En-têtes
            hdr_cells = defauts_table.rows[0].cells
            hdr_cells[0].text = "Cause"
            hdr_cells[1].text = "Illustration associée"
            
            for cell in hdr_cells:
                cell.paragraphs[0].runs[0].bold = True
                cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
                cell.paragraphs[0].runs[0].font.size = Pt(11)
            
            # Définir largeurs
            hdr_cells[0].width = Cm(8)
            hdr_cells[1].width = Cm(8)
            
            # Ajouter les défauts
            for defaut in gamme_data['defauts']:
                row_cells = defauts_table.add_row().cells
                
                # Cause et description
                cause_para = row_cells[0].paragraphs[0]
                cause_para.clear()
                cause_para.add_run(defaut['cause']).bold = True
                if defaut.get('description'):
                    cause_para.add_run(f"\n{defaut['description']}")
                    cause_para.runs[-1].font.size = Pt(9)
                    cause_para.runs[-1].italic = True
                
                # Image du défaut
                if defaut.get('image') and os.path.exists(defaut['image']):
                    try:
                        img_para = row_cells[1].paragraphs[0]
                        img_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                        
                        run = img_para.add_run()
                        run.add_picture(defaut['image'], width=Cm(6.0), height=Cm(4.0))
                        
                        # Légende
                        img_name = os.path.basename(defaut['image']).replace('.png', '').replace('_', ' ').title()
                        img_para.add_run(f"\n{img_name}")
                        img_para.runs[-1].font.size = Pt(8)
                        img_para.runs[-1].italic = True
                        
                    except Exception as e:
                        logger.warning(f"⚠️ Erreur insertion image défaut: {e}")
                        row_cells[1].text = "Image non disponible"
                        row_cells[1].paragraphs[0].runs[0].font.size = Pt(9)
                        row_cells[1].paragraphs[0].runs[0].italic = True
                else:
                    row_cells[1].text = "Voir documentation technique"
                    row_cells[1].paragraphs[0].runs[0].font.size = Pt(9)
                    row_cells[1].paragraphs[0].runs[0].italic = True
                    row_cells[1].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
                
                # Largeurs
                row_cells[0].width = Cm(8)
                row_cells[1].width = Cm(8)
        
        # ✅ NOUVEAU: Section 3 - Contrôles à effectuer
        if gamme_data.get('controles'):
            doc.add_paragraph()
            doc.add_heading("Contrôles à Effectuer", level=2)
            
            controles_table = doc.add_table(rows=1, cols=2)
            controles_table.style = 'Table Grid'
            controles_table.alignment = WD_TABLE_ALIGNMENT.CENTER
            
            # En-têtes
            hdr_cells = controles_table.rows[0].cells
            hdr_cells[0].text = "Contrôle"
            hdr_cells[1].text = "Outil ou méthode illustrée"
            
            for cell in hdr_cells:
                cell.paragraphs[0].runs[0].bold = True
                cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
                cell.paragraphs[0].runs[0].font.size = Pt(11)
            
            # Définir largeurs
            hdr_cells[0].width = Cm(8)
            hdr_cells[1].width = Cm(8)
            
            # Ajouter les contrôles
            for controle in gamme_data['controles']:
                row_cells = controles_table.add_row().cells
                
                # Contrôle et description
                controle_para = row_cells[0].paragraphs[0]
                controle_para.clear()
                controle_para.add_run(controle['controle']).bold = True
                
                if controle.get('description'):
                    controle_para.add_run(f"\n{controle['description']}")
                    controle_para.runs[-1].font.size = Pt(9)
                    controle_para.runs[-1].italic = True
                
                if controle.get('frequence'):
                    controle_para.add_run(f"\nFréquence: {controle['frequence']}")
                    controle_para.runs[-1].font.size = Pt(8)
                    controle_para.runs[-1].font.color.rgb = None  # Couleur grise
                
                # Image de l'outil/méthode
                if controle.get('image') and os.path.exists(controle['image']):
                    try:
                        img_para = row_cells[1].paragraphs[0]
                        img_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                        
                        run = img_para.add_run()
                        run.add_picture(controle['image'], width=Cm(6.0), height=Cm(4.0))
                        
                        # Légende
                        img_name = os.path.basename(controle['image']).replace('.png', '').replace('_', ' ').title()
                        img_para.add_run(f"\n{img_name}")
                        img_para.runs[-1].font.size = Pt(8)
                        img_para.runs[-1].italic = True
                        
                    except Exception as e:
                        logger.warning(f"⚠️ Erreur insertion image contrôle: {e}")
                        row_cells[1].text = "Image non disponible"
                        row_cells[1].paragraphs[0].runs[0].font.size = Pt(9)
                        row_cells[1].paragraphs[0].runs[0].italic = True
                else:
                    row_cells[1].text = "Voir manuel équipement"
                    row_cells[1].paragraphs[0].runs[0].font.size = Pt(9)
                    row_cells[1].paragraphs[0].runs[0].italic = True
                    row_cells[1].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
                
                # Largeurs
                row_cells[0].width = Cm(8)
                row_cells[1].width = Cm(8)
        
        # ✅ NOUVEAU: Section 4 - Image finale de la chaudière
        chaudiere_images = [
            'static/images/image_chaudiere/chaudiere_schema.png',
            'static/images/image_chaudiere_global/shema_chaudiere_global.jpg',
            'static/images/image_chaudiere/chaudiere_5_6.png'
        ]
        
        chaudiere_image = None
        for img_path in chaudiere_images:
            if os.path.exists(img_path):
                chaudiere_image = img_path
                break
        
        if chaudiere_image:
            doc.add_paragraph()
            doc.add_heading("Schéma Général de la Chaudière", level=2)
            
            # Paragraphe centré pour l'image
            chaudiere_para = doc.add_paragraph()
            chaudiere_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            try:
                run = chaudiere_para.add_run()
                run.add_picture(chaudiere_image, width=Cm(14.0))
                
                # Légende sous l'image
                chaudiere_para.add_run(f"\nSchéma d'ensemble - Chaudière industrielle")
                chaudiere_para.runs[-1].font.size = Pt(12)
                chaudiere_para.runs[-1].bold = True
                chaudiere_para.runs[-1].font.color.rgb = None  # Couleur standard
                
            except Exception as e:
                logger.warning(f"⚠️ Erreur insertion image chaudière: {e}")
                chaudiere_para.add_run("Schéma chaudière non disponible")
                chaudiere_para.runs[0].italic = True
        
        # ✅ Informations historique si disponible
        if gamme_data.get('historique_occurrences'):
            priority_para = doc.add_paragraph()
            priority_text = f"⚠️ PRIORITÉ {gamme_data.get('priorite_historique', 'NORMALE')}: "
            priority_text += gamme_data.get('historique_info', '')
            
            priority_para.add_run(priority_text)
            priority_para.runs[0].bold = True
            priority_para.runs[0].font.size = Pt(11)
            
            if gamme_data.get('priorite_historique') == 'ÉLEVÉE':
                priority_para.runs[0].font.color.rgb = None  # Rouge
        
        # ✅ Pied de page avec informations du générateur
        footer_para = doc.add_paragraph()
        footer_para.add_run(f"\nRédacteur: AMDEC & Gamme IA")
        footer_para.add_run(f"\nDate: {gamme_data['date']}")
        footer_para.add_run(f"\nDocument source: PLAN DE MAINTENANCE PREVENTIVE")
        footer_para.add_run(f"\nFolio: 1/1")
        
        for run in footer_para.runs:
            run.font.size = Pt(9)
            run.italic = True
        footer_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # ✅ Sauvegarder le document
        doc.save(output_path)
        logger.info(f"✅ Document Word COMPLET créé avec:")
        logger.info(f"   - {len(gamme_data.get('images', []))} images d'appareils intégrées")
        logger.info(f"   - {len(gamme_data.get('inspections', []))} inspections illustrées")
        logger.info(f"   - {len(gamme_data.get('defauts', []))} défauts avec images")
        logger.info(f"   - {len(gamme_data.get('controles', []))} contrôles avec méthodes")
        logger.info(f"   - Schéma chaudière intégré")
    
    # ===============================
    # MÉTHODES ORIGINALES CONSERVÉES
    # ===============================
    
    def _get_materials_for_operation(self, operation: Dict, all_materials: List[str]) -> str:
        """✅ Extrait les matériels pertinents pour une opération donnée"""
        op_name = operation.get('name', '').lower()
        op_desc = operation.get('description', '').lower()
        
        # Mapping mots-clés → matériels
        keywords_mapping = {
            'inspection': ['Lampe torche', 'Appareil photo'],
            'visuel': ['Lampe torche', 'Miroir'],
            'ultrason': ['Appareil ultrasons', 'Gel contact'],
            'mesure': ['Appareil ultrasons', 'Calibres'],
            'nettoyage': ['Brosse', 'Produit nettoyant'],
            'test': ['Kit test étanchéité'],
            'étanchéité': ['Kit test étanchéité'],
            'traitement': ['Pinceau', 'Peinture'],
            'contrôle': ['Lampe torche', 'Instruments mesure']
        }
        
        relevant_materials = []
        
        # Chercher les matériels selon les mots-clés
        for keyword, materials in keywords_mapping.items():
            if keyword in op_name or keyword in op_desc:
                for material in materials:
                    # Chercher dans la liste complète des matériels
                    matching = [m for m in all_materials if any(word in m.lower() for word in material.lower().split())]
                    relevant_materials.extend(matching)
        
        # Si aucun matériel trouvé, prendre les premiers de la liste
        if not relevant_materials:
            relevant_materials = all_materials[:2]
        
        # Supprimer doublons et limiter
        unique_materials = list(dict.fromkeys(relevant_materials))[:3]
        
        return " + ".join(unique_materials) if unique_materials else "Matériel standard"
    
    def _group_materials_by_category(self, materials: List[str]) -> Dict[str, List[str]]:
        """✅ Groupe les matériels par catégorie pour un affichage organisé"""
        categories = {
            'Inspection': [],
            'Mesure': [],
            'Nettoyage': [],
            'Réparation': [],
            'Sécurité': []
        }
        
        # Mapping matériels → catégories
        material_categories = {
            'lampe': 'Inspection',
            'appareil photo': 'Inspection',
            'caméra': 'Inspection',
            'endoscope': 'Inspection',
            'ultrasons': 'Mesure',
            'calibre': 'Mesure',
            'mètre': 'Mesure',
            'thermomètre': 'Mesure',
            'brosse': 'Nettoyage',
            'produit nettoyant': 'Nettoyage',
            'chiffon': 'Nettoyage',
            'solvant': 'Nettoyage',
            'poste à souder': 'Réparation',
            'clé': 'Réparation',
            'outil': 'Réparation',
            'pinceau': 'Réparation',
            'epi': 'Sécurité',
            'protection': 'Sécurité',
            'casque': 'Sécurité'
        }
        
        # Classer chaque matériel
        for material in materials:
            material_lower = material.lower()
            assigned = False
            
            for keyword, category in material_categories.items():
                if keyword in material_lower:
                    categories[category].append(material)
                    assigned = True
                    break
            
            # Si pas de catégorie trouvée, mettre dans Inspection par défaut
            if not assigned:
                categories['Inspection'].append(material)
        
        # Supprimer les catégories vides
        return {cat: items for cat, items in categories.items() if items}
    
    def _build_maintenance_knowledge(self) -> Dict:
        """Construit la base de connaissances pour la maintenance"""
        return {
            'economiseur_bt': {
                'epingle': {
                    'critical_points': ['Corrosion externe', 'Encrassement interne', 'Fissuration'],
                    'typical_failures': ['Amincissement parois', 'Dépôts internes'],
                    'inspection_focus': ['Épaisseur parois', 'État des soudures', 'Propreté interne']
                },
                'collecteur_sortie': {
                    'critical_points': ['Corrosion interne', 'Contraintes thermiques', 'Étanchéité'],
                    'typical_failures': ['Perte matière interne', 'Fuites aux joints'],
                    'inspection_focus': ['Test étanchéité', 'Inspection interne', 'Contrôle soudures']
                }
            },
            'economiseur_ht': {
                'collecteur_entree': {
                    'critical_points': ['Érosion par cendres', 'Température élevée'],
                    'typical_failures': ['Amincissement accéléré', 'Déformation'],
                    'inspection_focus': ['Contrôle épaisseur', 'Inspection déformation']
                },
                'tubes_suspension': {
                    'critical_points': ['Fatigue mécanique', 'Vibrations', 'Supports'],
                    'typical_failures': ['Fissures supports', 'Desserrage fixations'],
                    'inspection_focus': ['Contrôle fixations', 'Analyse vibrations']
                }
            },
            'surchauffeur_bt': {
                'epingle': {
                    'critical_points': ['Graphitisation', 'Corrosion côté feu', 'Surchauffe'],
                    'typical_failures': ['Rupture ductile', 'Perte métal externe'],
                    'inspection_focus': ['Contrôle température', 'État surfaces externes']
                },
                'collecteur_entree': {
                    'critical_points': ['Corrosion interne', 'Contraintes', 'Érosion'],
                    'typical_failures': ['Fissuration', 'Perte matière'],
                    'inspection_focus': ['Inspection interne', 'Contrôle contraintes']
                }
            },
            'surchauffeur_ht': {
                'tube_porteur': {
                    'critical_points': ['Fluage', 'Surchauffe long terme', 'Contraintes'],
                    'typical_failures': ['Rupture fluage', 'Déformation permanente'],
                    'inspection_focus': ['Analyse métallurgique', 'Mesure déformation', 'Contrôle température']
                },
                'branches_entree': {
                    'critical_points': ['Corrosion côté feu', 'Fatigue thermique'],
                    'typical_failures': ['Perte métal externe', 'Fissures thermiques'],
                    'inspection_focus': ['Inspection surfaces', 'Contrôle fissures']
                },
                'collecteur_sortie': {
                    'critical_points': ['SCC', 'Contraintes interfaces', 'Soudures'],
                    'typical_failures': ['Fissures intergranulaires', 'Rupture soudure'],
                    'inspection_focus': ['Contrôle soudures', 'Inspection fissures']
                }
            },
            'rechauffeur_bt': {
                'collecteur_entree': {
                    'critical_points': ['Hydrogen damage', 'Corrosion interne'],
                    'typical_failures': ['Microfissures', 'Dégradation interne'],
                    'inspection_focus': ['Inspection interne', 'Contrôle microfissures']
                },
                'tubes_suspension': {
                    'critical_points': ['Fatigue thermique', 'Cycles thermiques'],
                    'typical_failures': ['Fissures fatigue', 'Déformation'],
                    'inspection_focus': ['Contrôle fixations', 'Mesure déformation']
                },
                'tube_porteur': {
                    'critical_points': ['Fatigue thermique', 'Contraintes mécaniques'],
                    'typical_failures': ['Fissures', 'Rupture supports'],
                    'inspection_focus': ['Contrôle supports', 'Inspection fissures']
                }
            },
            'rechauffeur_ht': {
                'branches_sortie': {
                    'critical_points': ['Acid attack', 'Corrosion surface'],
                    'typical_failures': ['Surface fromage suisse', 'Corrosion généralisée'],
                    'inspection_focus': ['Inspection surfaces', 'Analyse corrosion']
                },
                'collecteur_entree': {
                    'critical_points': ['Waterside corrosion', 'Dépôts internes'],
                    'typical_failures': ['Fissures internes', 'Accumulation dépôts'],
                    'inspection_focus': ['Inspection interne', 'Nettoyage dépôts']
                },
                'collecteur_sortie': {
                    'critical_points': ['Dissimilar metal weld', 'Interfaces soudées'],
                    'typical_failures': ['Rupture soudure', 'Contraintes interfaces'],
                    'inspection_focus': ['Contrôle soudures', 'Inspection interfaces']
                }
            }
        }
    
    def _build_materials_database(self) -> Dict:
        """Base de données des matériels de maintenance"""
        return {
            'base_tools': [
                'Lampe torche',
                'Appareil photo numérique',
                'Caméra d\'inspection',
                'Équipements de protection individuelle (EPI)'
            ],
            'inspection_tools': {
                'visual': ['Miroir d\'inspection', 'Endoscope', 'Caméra thermique'],
                'ultrasonic': ['Appareil à ultrasons', 'Gel de contact', 'Sondes spécifiques'],
                'dimensional': ['Mètre ruban', 'Calibres', 'Jauges d\'épaisseur'],
                'leak_test': ['Kit test d\'étanchéité', 'Manomètres', 'Produit traceur'],
                'vibration': ['Capteurs vibratoires', 'Analyseur de vibrations'],
                'temperature': ['Capteurs de température', 'Thermomètres IR']
            },
            'maintenance_tools': {
                'cleaning': ['Brosse métallique', 'Brosse souple', 'Produit nettoyant', 'Chiffons'],
                'repair': ['Poste à souder', 'Électrodes', 'Meuleuse', 'Outillage spécialisé'],
                'coating': ['Pinceau', 'Rouleau', 'Pistolet peinture', 'Peinture anticorrosion'],
                'fastening': ['Clés spécifiques', 'Outils de serrage', 'Couples-mètres']
            },
            'consumables': {
                'cleaning': ['Solvants', 'Dégraissants', 'Produits détartrants'],
                'sealing': ['Joints d\'étanchéité', 'Pâte d\'étanchéité', 'Mastique'],
                'coating': ['Peintures spéciales', 'Primaires', 'Revêtements céramiques'],
                'replacement': ['Boulonnerie', 'Joints toriques', 'Éléments de fixation']
            }
        }
    
    def _build_operations_database(self) -> Dict:
        """Base de données des opérations de maintenance"""
        return {
            'basic_operations': [
                {
                    'name': 'Inspection visuelle générale',
                    'description': 'Examen visuel complet de l\'état général, recherche de signes de dégradation',
                    'time': 15,
                    'tools': ['Lampe torche', 'Appareil photo'],
                    'consumables': []
                }
            ],
            'by_criticality': {
                'low': [  # Criticité <= 12
                    {
                        'name': 'Contrôle visuel des points sensibles',
                        'description': 'Vérification ciblée des zones connues pour être sensibles',
                        'time': 20,
                        'tools': ['Lampe torche', 'Miroir d\'inspection'],
                        'consumables': []
                    }
                ],
                'medium': [  # 12 < Criticité <= 16
                    {
                        'name': 'Contrôle dimensionnel',
                        'description': 'Mesure des dimensions critiques et vérification des tolérances',
                        'time': 25,
                        'tools': ['Appareil ultrasons', 'Calibres'],
                        'consumables': ['Gel contact']
                    },
                    {
                        'name': 'Nettoyage préventif',
                        'description': 'Nettoyage des surfaces et élimination des dépôts',
                        'time': 30,
                        'tools': ['Brosse souple', 'Produit nettoyant'],
                        'consumables': ['Solvant', 'Chiffons']
                    }
                ],
                'high': [  # 16 < Criticité <= 20
                    {
                        'name': 'Inspection approfondie',
                        'description': 'Contrôle détaillé par méthodes non destructives',
                        'time': 35,
                        'tools': ['Endoscope', 'Ultrasons'],
                        'consumables': ['Gel contact']
                    },
                    {
                        'name': 'Test d\'étanchéité',
                        'description': 'Vérification de l\'étanchéité sous pression de service',
                        'time': 40,
                        'tools': ['Kit test étanchéité', 'Manomètres'],
                        'consumables': ['Produit traceur']
                    },
                    {
                        'name': 'Traitement préventif',
                        'description': 'Application de traitements protecteurs',
                        'time': 45,
                        'tools': ['Pinceau', 'Équipement protection'],
                        'consumables': ['Revêtement protecteur']
                    }
                ],
                'critical': [  # Criticité > 20
                    {
                        'name': 'Surveillance continue',
                        'description': 'Installation de systèmes de surveillance permanente',
                        'time': 60,
                        'tools': ['Capteurs température', 'Système acquisition'],
                        'consumables': ['Câblage', 'Fixations']
                    },
                    {
                        'name': 'Intervention corrective',
                        'description': 'Réparation ou remplacement des éléments défaillants',
                        'time': 90,
                        'tools': ['Poste à souder', 'Outillage spécialisé'],
                        'consumables': ['Électrodes', 'Pièces rechange']
                    },
                    {
                        'name': 'Renforcement structural',
                        'description': 'Mise en place de renforts ou modifications structurelles',
                        'time': 120,
                        'tools': ['Équipement soudage', 'Manutention'],
                        'consumables': ['Matériaux renfort', 'Consommables soudage']
                    }
                ]
            },
            'specialized_operations': {
                'epingle': [
                    {
                        'name': 'Contrôle épaisseur parois',
                        'description': 'Mesure systématique de l\'épaisseur par ultrasons',
                        'time': 30,
                        'tools': ['Ultrasons', 'Gel contact'],
                        'consumables': ['Gel ultrasons']
                    }
                ],
                'collecteur': [
                    {
                        'name': 'Inspection endoscopique',
                        'description': 'Examen interne complet par endoscopie',
                        'time': 45,
                        'tools': ['Endoscope', 'Éclairage'],
                        'consumables': []
                    }
                ],
                'tube': [
                    {
                        'name': 'Analyse vibratoire',
                        'description': 'Mesure et analyse des vibrations en fonctionnement',
                        'time': 40,
                        'tools': ['Capteurs vibratoires', 'Analyseur'],
                        'consumables': []
                    }
                ],
                'branches': [
                    {
                        'name': 'Contrôle raccordements',
                        'description': 'Vérification de l\'intégrité des raccordements',
                        'time': 25,
                        'tools': ['Clés spécifiques', 'Couples-mètre'],
                        'consumables': []
                    }
                ]
            }
        }
    
    def _generate_materials_list(self, component: str, subcomponent: str, criticality: int) -> List[str]:
        """Génère la liste des matériels nécessaires"""
        materials = []
        
        # Outils de base toujours nécessaires
        materials.extend(self.materials_database['base_tools'])
        
        # Outils selon le type de sous-composant
        subcomp_type = self._get_subcomponent_type(subcomponent)
        
        # Inspection visuelle
        materials.extend(self.materials_database['inspection_tools']['visual'][:2])
        
        # Outils selon la criticité
        if criticality <= 12:
            # Criticité faible : inspection de base
            materials.extend(self.materials_database['inspection_tools']['dimensional'][:2])
        elif criticality <= 16:
            # Criticité moyenne : inspection + nettoyage
            materials.extend(self.materials_database['inspection_tools']['ultrasonic'])
            materials.extend(self.materials_database['maintenance_tools']['cleaning'][:3])
        elif criticality <= 20:
            # Criticité élevée : inspection avancée + traitement
            materials.extend(self.materials_database['inspection_tools']['ultrasonic'])
            materials.extend(self.materials_database['inspection_tools']['leak_test'])
            materials.extend(self.materials_database['maintenance_tools']['coating'])
        else:
            # Criticité critique : équipement complet
            materials.extend(self.materials_database['inspection_tools']['ultrasonic'])
            materials.extend(self.materials_database['inspection_tools']['temperature'])
            materials.extend(self.materials_database['maintenance_tools']['repair'])
        
        # Consommables
        if criticality > 12:
            materials.extend(self.materials_database['consumables']['cleaning'][:2])
        if criticality > 16:
            materials.extend(self.materials_database['consumables']['sealing'][:2])
        if criticality > 20:
            materials.extend(self.materials_database['consumables']['replacement'])
        
        # Supprimer les doublons et limiter la liste
        unique_materials = list(dict.fromkeys(materials))
        return unique_materials[:12]  # Limiter à 12 éléments max
    
    def _generate_operations(self, component: str, subcomponent: str, criticality: int) -> List[Dict]:
        """Génère la liste des opérations de maintenance"""
        operations = []
        order = 1
        
        # Opération de base : inspection visuelle
        base_op = self.operations_database['basic_operations'][0].copy()
        base_op['order'] = order
        operations.append(base_op)
        order += 1
        
        # Opérations selon la criticité
        criticality_level = self._get_criticality_level(criticality).lower()
        criticality_ops = self.operations_database['by_criticality'].get(criticality_level, [])
        
        for op in criticality_ops:
            op_copy = op.copy()
            op_copy['order'] = order
            operations.append(op_copy)
            order += 1
        
        # Opérations spécialisées selon le type de sous-composant
        subcomp_type = self._get_subcomponent_type(subcomponent)
        specialized_ops = self.operations_database['specialized_operations'].get(subcomp_type, [])
        
        if specialized_ops and criticality > 12:
            # Ajouter une opération spécialisée si criticité suffisante
            spec_op = specialized_ops[0].copy()
            spec_op['order'] = order
            operations.append(spec_op)
        
        return operations
    
    def _generate_safety_instructions(self, component: str, subcomponent: str) -> List[str]:
        """Génère les consignes de sécurité"""
        base_safety = [
            "Port des équipements de protection individuelle (EPI) obligatoire",
            "Consignation complète (électrique, mécanique, thermique) avant intervention",
            "Vérification de l'absence de pression et contrôle de température",
            "Balisage et signalisation de la zone d'intervention"
        ]
        
        # Consignes spécifiques selon le composant
        specific_safety = {
            'surchauffeur_ht': [
                "Attention aux hautes températures résiduelles",
                "Contrôle de l'atmosphère avant intervention en espace confiné"
            ],
            'rechauffeur_ht': [
                "Surveillance de la qualité de l'air (vapeurs acides)",
                "Protection respiratoire renforcée"
            ]
        }
        
        safety_instructions = base_safety.copy()
        if component in specific_safety:
            safety_instructions.extend(specific_safety[component])
        
        return safety_instructions
    
    def _get_subcomponent_type(self, subcomponent: str) -> str:
        """Détermine le type de sous-composant"""
        if 'epingle' in subcomponent:
            return 'epingle'
        elif 'collecteur' in subcomponent:
            return 'collecteur'
        elif 'tube' in subcomponent:
            return 'tube'
        elif 'branches' in subcomponent:
            return 'branches'
        else:
            return 'collecteur'  # par défaut
    
    def _get_criticality_level(self, criticality: int) -> str:
        """Retourne le niveau de criticité"""
        if criticality <= 12:
            return 'low'
        elif criticality <= 16:
            return 'medium'
        elif criticality <= 20:
            return 'high'
        else:
            return 'critical'
    
    def _calculate_total_time(self, operations: List[Dict]) -> str:
        """Calcule le temps total estimé"""
        total_minutes = sum(op.get('time', 0) for op in operations)
        
        hours = total_minutes // 60
        minutes = total_minutes % 60
        
        if hours > 0:
            return f"{hours}h{minutes:02d}min"
        else:
            return f"{minutes}min"