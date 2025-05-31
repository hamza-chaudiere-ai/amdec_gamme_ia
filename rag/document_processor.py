#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
Processeur de Documents pour AMDEC & Gamme IA
Extraction et traitement du contenu des documents techniques
"""

import os
import re
import logging
from typing import List, Dict, Tuple, Optional
from pathlib import Path
import pandas as pd
from docx import Document
import PyPDF2
from bs4 import BeautifulSoup
import markdown

logger = logging.getLogger(__name__)

class DocumentProcessor:
    """Processeur de documents techniques pour la base de connaissances"""
    
    def __init__(self, documents_dir: str = "data/documents"):
        """
        Initialise le processeur de documents
        
        Args:
            documents_dir: Répertoire contenant les documents à traiter
        """
        self.documents_dir = documents_dir
        self.supported_extensions = {
            '.docx': self._extract_from_docx,
            '.pdf': self._extract_from_pdf,
            '.xlsx': self._extract_from_excel,
            '.md': self._extract_from_markdown,
            '.txt': self._extract_from_text
        }
        
        # Créer le répertoire s'il n'existe pas
        os.makedirs(documents_dir, exist_ok=True)
    
    def process_all_documents(self) -> List[Dict]:
        """
        Traite tous les documents du répertoire
        
        Returns:
            Liste de dictionnaires contenant les extraits de documents
        """
        documents = []
        
        try:
            if not os.path.exists(self.documents_dir):
                logger.warning(f"Répertoire {self.documents_dir} non trouvé")
                return self._create_default_knowledge_base()
            
            # Parcourir tous les fichiers
            for root, dirs, files in os.walk(self.documents_dir):
                for file in files:
                    file_path = os.path.join(root, file)
                    file_extension = Path(file).suffix.lower()
                    
                    if file_extension in self.supported_extensions:
                        logger.info(f"Traitement du document: {file}")
                        
                        try:
                            doc_content = self.supported_extensions[file_extension](file_path)
                            if doc_content:
                                documents.extend(doc_content)
                                
                        except Exception as e:
                            logger.error(f"Erreur traitement {file}: {e}")
                            continue
            
            # Si aucun document trouvé, créer une base par défaut
            if not documents:
                logger.info("Aucun document trouvé, création de la base de connaissances par défaut")
                documents = self._create_default_knowledge_base()
            
            logger.info(f"Total: {len(documents)} extraits de documents traités")
            return documents
            
        except Exception as e:
            logger.error(f"Erreur lors du traitement des documents: {e}")
            return self._create_default_knowledge_base()
    
    def _extract_from_docx(self, file_path: str) -> List[Dict]:
        """Extrait le contenu d'un fichier Word"""
        try:
            doc = Document(file_path)
            extracts = []
            
            filename = os.path.basename(file_path)
            
            # Extraire par paragraphes
            current_section = ""
            content_buffer = []
            
            for para in doc.paragraphs:
                text = para.text.strip()
                
                if not text:
                    continue
                
                # Détecter les titres (texte en gras ou style titre)
                is_heading = (
                    para.style.name.startswith('Heading') or
                    (para.runs and para.runs[0].bold and len(text) < 100)
                )
                
                if is_heading:
                    # Sauvegarder la section précédente
                    if content_buffer:
                        extracts.append({
                            'content': '\n'.join(content_buffer),
                            'source': filename,
                            'section': current_section,
                            'type': 'document_section'
                        })
                        content_buffer = []
                    
                    current_section = text
                else:
                    content_buffer.append(text)
            
            # Sauvegarder la dernière section
            if content_buffer:
                extracts.append({
                    'content': '\n'.join(content_buffer),
                    'source': filename,
                    'section': current_section,
                    'type': 'document_section'
                })
            
            # Extraire les tableaux
            for i, table in enumerate(doc.tables):
                table_content = self._extract_table_content(table)
                if table_content:
                    extracts.append({
                        'content': table_content,
                        'source': filename,
                        'section': f'Tableau {i+1}',
                        'type': 'table'
                    })
            
            logger.info(f"Document Word {filename}: {len(extracts)} extraits")
            return extracts
            
        except Exception as e:
            logger.error(f"Erreur extraction Word {file_path}: {e}")
            return []
    
    def _extract_from_pdf(self, file_path: str) -> List[Dict]:
        """Extrait le contenu d'un fichier PDF"""
        try:
            extracts = []
            filename = os.path.basename(file_path)
            
            with open(file_path, 'rb') as file:
                pdf_reader = PyPDF2.PdfReader(file)
                
                for page_num, page in enumerate(pdf_reader.pages):
                    text = page.extract_text()
                    
                    if text.strip():
                        # Diviser en paragraphes
                        paragraphs = [p.strip() for p in text.split('\n\n') if p.strip()]
                        
                        for para in paragraphs:
                            if len(para) > 50:  # Ignorer les textes trop courts
                                extracts.append({
                                    'content': para,
                                    'source': filename,
                                    'section': f'Page {page_num + 1}',
                                    'type': 'pdf_paragraph'
                                })
            
            logger.info(f"Document PDF {filename}: {len(extracts)} extraits")
            return extracts
            
        except Exception as e:
            logger.error(f"Erreur extraction PDF {file_path}: {e}")
            return []
    
    def _extract_from_excel(self, file_path: str) -> List[Dict]:
        """Extrait le contenu d'un fichier Excel"""
        try:
            extracts = []
            filename = os.path.basename(file_path)
            
            # Lire toutes les feuilles
            excel_file = pd.ExcelFile(file_path)
            
            for sheet_name in excel_file.sheet_names:
                df = pd.read_excel(file_path, sheet_name=sheet_name)
                
                if df.empty:
                    continue
                
                # Convertir le DataFrame en texte structuré
                sheet_content = f"Feuille: {sheet_name}\n"
                sheet_content += f"Colonnes: {', '.join(df.columns)}\n\n"
                
                # Ajouter quelques lignes d'exemple
                for idx, row in df.head(10).iterrows():
                    row_text = " | ".join([f"{col}: {val}" for col, val in row.items() if pd.notna(val)])
                    sheet_content += f"Ligne {idx + 1}: {row_text}\n"
                
                extracts.append({
                    'content': sheet_content,
                    'source': filename,
                    'section': sheet_name,
                    'type': 'excel_sheet'
                })
            
            logger.info(f"Document Excel {filename}: {len(extracts)} extraits")
            return extracts
            
        except Exception as e:
            logger.error(f"Erreur extraction Excel {file_path}: {e}")
            return []
    
    def _extract_from_markdown(self, file_path: str) -> List[Dict]:
        """Extrait le contenu d'un fichier Markdown"""
        try:
            with open(file_path, 'r', encoding='utf-8') as f:
                content = f.read()
            
            # Convertir en HTML puis extraire
            html = markdown.markdown(content)
            soup = BeautifulSoup(html, 'html.parser')
            
            extracts = []
            filename = os.path.basename(file_path)
            
            # Extraire par sections (titres)
            current_section = ""
            content_buffer = []
            
            for element in soup.find_all(['h1', 'h2', 'h3', 'h4', 'h5', 'h6', 'p', 'ul', 'ol']):
                if element.name.startswith('h'):
                    # Nouveau titre
                    if content_buffer:
                        extracts.append({
                            'content': '\n'.join(content_buffer),
                            'source': filename,
                            'section': current_section,
                            'type': 'markdown_section'
                        })
                        content_buffer = []
                    
                    current_section = element.get_text().strip()
                else:
                    content_buffer.append(element.get_text().strip())
            
            # Dernière section
            if content_buffer:
                extracts.append({
                    'content': '\n'.join(content_buffer),
                    'source': filename,
                    'section': current_section,
                    'type': 'markdown_section'
                })
            
            return extracts
            
        except Exception as e:
            logger.error(f"Erreur extraction Markdown {file_path}: {e}")
            return []
    
    def _extract_from_text(self, file_path: str) -> List[Dict]:
        """Extrait le contenu d'un fichier texte"""
        try:
            with open(file_path, 'r', encoding='utf-8') as f:
                content = f.read()
            
            # Diviser en paragraphes
            paragraphs = [p.strip() for p in content.split('\n\n') if p.strip()]
            
            extracts = []
            filename = os.path.basename(file_path)
            
            for i, para in enumerate(paragraphs):
                if len(para) > 30:  # Ignorer les paragraphes trop courts
                    extracts.append({
                        'content': para,
                        'source': filename,
                        'section': f'Paragraphe {i+1}',
                        'type': 'text_paragraph'
                    })
            
            return extracts
            
        except Exception as e:
            logger.error(f"Erreur extraction texte {file_path}: {e}")
            return []
    
    def _extract_table_content(self, table) -> str:
        """Extrait le contenu d'un tableau Word"""
        try:
            content = []
            
            for row in table.rows:
                row_content = []
                for cell in row.cells:
                    cell_text = cell.text.strip()
                    if cell_text:
                        row_content.append(cell_text)
                
                if row_content:
                    content.append(" | ".join(row_content))
            
            return "\n".join(content) if content else ""
            
        except Exception as e:
            logger.error(f"Erreur extraction tableau: {e}")
            return ""
    
    def _create_default_knowledge_base(self) -> List[Dict]:
        """Crée une base de connaissances par défaut basée sur l'expertise du projet"""
        logger.info("Création de la base de connaissances par défaut")
        
        knowledge_base = []
        
        # Base de connaissances AMDEC
        amdec_knowledge = {
            "Économiseur BT - Épingle": {
                "description": "L'épingle de l'économiseur basse température est un élément crucial pour le transfert thermique. Elle est sujette à la corrosion externe et à l'encrassement interne.",
                "defauts_courants": [
                    "Corrosion externe par les fumées",
                    "Encrassement interne réduisant l'efficacité",
                    "Fissuration due aux cycles thermiques",
                    "Érosion par les particules de cendres"
                ],
                "criticite": "24 (F=3, G=4, D=2)",
                "solutions": [
                    "Revêtement céramique pour protection contre la corrosion",
                    "Contrôle de la qualité de l'eau d'alimentation",
                    "Nettoyage régulier par ultrasons",
                    "Surveillance de la température de fonctionnement"
                ],
                "maintenance_preventive": "Inspection trimestrielle avec contrôle par ultrasons"
            },
            
            "Économiseur BT - Collecteur sortie": {
                "description": "Le collecteur de sortie de l'économiseur BT collecte la vapeur produite. Point critique pour l'étanchéité et les contraintes thermiques.",
                "defauts_courants": [
                    "Caustic attack (attaque alcaline)",
                    "Fissuration par contraintes thermiques",
                    "Corrosion interne",
                    "Fuites aux joints et soudures"
                ],
                "criticite": "45 (F=3, G=5, D=3)",
                "solutions": [
                    "Traitement de l'eau pour contrôler la chimie",
                    "Inspection endoscopique régulière",
                    "Réparation par soudage qualifié",
                    "Remplacement des joints d'étanchéité"
                ],
                "maintenance_preventive": "Inspection mensuelle et test d'étanchéité"
            },
            
            "Surchauffeur HT - Tube porteur": {
                "description": "Le tube porteur du surchauffeur haute température supporte les contraintes mécaniques et thermiques les plus élevées. Élément critique pour la sécurité.",
                "defauts_courants": [
                    "Long-term overheat (surchauffe long terme)",
                    "Rupture par fluage",
                    "Corrosion haute température",
                    "Déformation permanente"
                ],
                "criticite": "30 (F=2, G=5, D=3)",
                "solutions": [
                    "Surveillance continue de la température",
                    "Analyse métallurgique périodique",
                    "Optimisation de la combustion",
                    "Installation d'alarmes de température"
                ],
                "maintenance_preventive": "Surveillance continue + inspection mensuelle"
            },
            
            "Réchauffeur HT - Branches sortie": {
                "description": "Les branches de sortie du réchauffeur HT sont exposées aux attaques acides. Zone particulièrement sensible à la corrosion.",
                "defauts_courants": [
                    "Acid attack (attaque acide)",
                    "Surface 'fromage suisse' par corrosion",
                    "Érosion par les vapeurs acides",
                    "Accumulation de dépôts acides"
                ],
                "criticite": "36 (F=3, G=4, D=3)",
                "solutions": [
                    "Protection chimique par injection d'additifs",
                    "Surveillance du pH et de la chimie",
                    "Revêtements résistants aux acides",
                    "Neutralisation des vapeurs acides"
                ],
                "maintenance_preventive": "Inspection trimestrielle avec analyse chimique"
            }
        }
        
        # Convertir en format de documents
        for component, data in amdec_knowledge.items():
            # Description générale
            knowledge_base.append({
                'content': f"""Composant: {component}
                
Description: {data['description']}

Criticité AMDEC: {data['criticite']}

Défauts courants identifiés:
{chr(10).join([f"• {defaut}" for defaut in data['defauts_courants']])}

Solutions techniques recommandées:
{chr(10).join([f"• {solution}" for solution in data['solutions']])}

Plan de maintenance préventive: {data['maintenance_preventive']}""",
                'source': 'Base_Connaissances_AMDEC',
                'section': component,
                'type': 'knowledge_base'
            })
        
        # Connaissances générales sur les défaillances
        general_knowledge = [
            {
                'content': """Types de corrosion dans les chaudières:

1. Corrosion externe (côté fumées):
   - Causée par les oxydes de soufre et chlorures
   - Température critique: 150-200°C
   - Solutions: revêtements céramiques, contrôle température

2. Corrosion interne (côté eau/vapeur):
   - Caustic attack: attaque alcaline à haute concentration
   - Acid attack: attaque acide par condensation
   - Solutions: traitement chimique de l'eau, surveillance pH

3. Corrosion haute température:
   - Oxydation accélérée au-delà de 550°C
   - Solutions: alliages résistants, surveillance température""",
                'source': 'Guide_Technique_Corrosion',
                'section': 'Types de corrosion',
                'type': 'technical_guide'
            },
            
            {
                'content': """Méthodes de contrôle non destructif (CND):

1. Contrôle par ultrasons:
   - Mesure d'épaisseur des parois
   - Détection de fissures internes
   - Fréquence recommandée: trimestrielle pour criticité élevée

2. Inspection endoscopique:
   - Examen visuel des surfaces internes
   - Détection de corrosion et dépôts
   - Utilisation pour les collecteurs et branches

3. Thermographie infrarouge:
   - Détection de points chauds
   - Surveillance des surchauffes locales
   - Contrôle en fonctionnement

4. Test d'étanchéité:
   - Vérification de l'intégrité des soudures
   - Pression d'épreuve selon normes
   - Obligatoire après réparation""",
                'source': 'Procedures_CND',
                'section': 'Méthodes de contrôle',
                'type': 'procedure'
            },
            
            {
                'content': """Actions correctives selon criticité AMDEC:

Criticité 1-12 (Négligeable):
- Maintenance corrective acceptable
- Surveillance annuelle
- Intervention à l'arrêt programmé

Criticité 13-16 (Moyenne):
- Maintenance préventive systématique
- Surveillance trimestrielle
- Intervention planifiée

Criticité 17-20 (Élevée):
- Maintenance préventive conditionnelle
- Surveillance mensuelle
- Intervention prioritaire

Criticité >20 (Critique):
- Remise en cause de la conception
- Surveillance continue
- Intervention immédiate si nécessaire""",
                'source': 'Matrice_Criticite_AMDEC',
                'section': 'Actions correctives',
                'type': 'methodology'
            }
        ]
        
        knowledge_base.extend(general_knowledge)
        
        logger.info(f"Base de connaissances par défaut créée: {len(knowledge_base)} entrées")
        return knowledge_base
    
    def preprocess_text(self, text: str) -> str:
        """Préprocess le texte pour améliorer la recherche vectorielle"""
        # Normaliser les espaces
        text = re.sub(r'\s+', ' ', text)
        
        # Supprimer les caractères spéciaux non utiles
        text = re.sub(r'[^\w\s\-.,;:()°%]', ' ', text)
        
        # Normaliser la casse pour les acronymes courants
        acronyms = {
            'amdec': 'AMDEC',
            'cnd': 'CND', 
            'epi': 'EPI',
            'bt': 'BT',
            'ht': 'HT'
        }
        
        for acronym, replacement in acronyms.items():
            text = re.sub(f'\\b{acronym}\\b', replacement, text, flags=re.IGNORECASE)
        
        return text.strip()